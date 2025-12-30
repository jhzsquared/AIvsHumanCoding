# Functions for cleaning PG and Peer healing data 
# create torch datasets from cleaned and pickled data

## Steps:
# 0. pre remove some weirdness (e.g. DW Complete has an instance where the D and I (interviewer) were mixed up)
# 1. Read in word doc by folder
# 2. Basic cleaning
#   - remove time stamps 
#   - remove fillers 
#       - (eg. headers, third person convo, (stopped), (pause), etc)
#   - separate out interviewer/interviewee
#   - Extract names (in case we want to normalize later)
# 3. Separate out chunks ("documents" that topics will be aligned with)
# 4. Save out in pickle files

# Resulting structure:
# 'file_name_clean.pickle': 
#   InterviewFile data type
#       List of InterviewLines + Metadata

# Notes: 
# - We don't remove speaker names.
#       Manually anonymized during data collection
# - Manually merged sentences in word docs that weren't assigned to a speaker

from dataclasses import dataclass, field
import unicodedata
import re
import docx
import torch
from torch.utils.data import Dataset
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pickle
import logging
import numpy as np
logger = logging.getLogger(__name__)

@dataclass
class InterviewLine():
    '''Data class for one line from an interview
    line: line number in original document
    raw_text: original text
    text: cleaned text
    interviewee/interviewer/credible_messenger: who is doing the speaking
        If unclear, then unknown is True
    useful_text: True if it's not an obvious weird header/blank line
    speaker: who the speaker of that line is (could be unknown)
    '''
    line: int 
    raw_text: str 
    text: str = field(default_factory = str) 
    interviewee: bool = False
    interviewer: bool = False
    credible_messenger: bool = False
    useful_text: bool = True
    speaker: str = field(default_factory = str, init = False)

@dataclass
class InterviewFile():
    '''Data class for a full transcription of an interview
    full_text: List of InterviewLines
    filename: originating filename
    focus_group: True if file is a transcription of a focus group
    interviewer_names: List of strings
    interviewee_names: List of strings
    start_line: what number line the interview text starts at (if header exists, this changes)
    end_line: what number the interview text ends at (sometimes empty lines exist)
    clean_lines: list of line numbers that are useful 
    '''
    
    full_text: list[InterviewLine]
    filename: str
    focus_group: bool 
    interviewer_names: list[str] = field(default_factory=list)
    interviewee_names: list[str] = field(default_factory=list)
    start_line: int = 0
    end_line: int = 0
    clean_lines: list[int] = field(default_factory=list)

    def __post_init__(self):
        # compile some final metadata items based on initiated data
        self.clean_lines = [line.line for line in self.full_text if line.useful_text]
        self.start_line = min(self.clean_lines)
        self.end_line = max(self.clean_lines)
        self.interviewee_names = set([line.speaker for line in self.full_text if line.interviewee ])
        self.interviewer_names = set([line.speaker for line in self.full_text if line.interviewer ])


def match_speakers(line: InterviewLine, focus_group = False, peer_healing = False):
    '''Identify the speaker in an InterviewLine
    Updates InterviewLine metadata with speaker name and if the interviewer/interviewee is speaking
    '''
    if focus_group:
        basic_speaker = re.compile(r'(?!Interviewer)(?!Drs*\.)(?!Interviewee)[0-9a-z|-|{Mr. }|{Ms. }| {Mrs. }]+(:|\.)', re.IGNORECASE)
    elif peer_healing:
        basic_speaker = re.compile(r'(?!Interviewer)(?!Drs*\.)(?!Interviewee)[0-9a-z|-]+(:|\.|-)', re.IGNORECASE)
    else:
        basic_speaker = re.compile(r'(?!Interviewer)(?!Drs*\.)(?!Interviewee)[0-9a-z|-]+(:|\.)', re.IGNORECASE)
    if peer_healing:
        interviewer = re.compile(r'I:|Host')
    else:
        interviewer = re.compile(r'(I)|(I2)|(Dr)|(1)|(Joe)', re.IGNORECASE)
    bad_words = re.compile('(stopped)|(pause)|(48)')
    speaker = re.match(basic_speaker, line.raw_text)
    if speaker:
        line.speaker = speaker[0]
        if re.match(bad_words, speaker[0]) is None:
            if re.match(interviewer, speaker[0]) is not None:
                line.interviewer = True
            else:
                line.interviewee =  True
    elif focus_group: # check if it's credible messenger
        cm = re.compile("Che:|Che")
        speaker = re.match(cm, line.raw_text)
        if speaker:
            line.speaker = speaker[0]
            line.credible_messenger = True
    else:
        line.useful_text = False # noise
        
    return line

def clean_line(line: InterviewLine):
    '''Basic cleaning 
    '''
    dirtytext = re.compile(r'(\(.+?\))|(\*\*)|(\[-*\])') # match text in (text), **, or [--]
    clean_text = re.sub(line.speaker, "", line.raw_text).strip()
    clean_text = re.sub(dirtytext, "", clean_text).strip()
    clean_text = clean_text.replace(u'\xa0', u' ')
    if clean_text == "":
        line.useful_text = False
    line.text = clean_text
    return line


def extract_file(filepath: str, focus_group: bool, peer_healing: bool):
    '''Extract raw interview data from file
    Parameters:
        filepath: path to raw docx file
        focus_group: whether or not file comes from a focus group
    '''
    doc = docx.Document(filepath)
    raw_text = [p.text.strip() for p in doc.paragraphs]
    full_text = [InterviewLine(line = i, 
                               raw_text = p)
                               for i,p in enumerate(raw_text)]
    if 'Wave 3' in filepath: #transcriptions are a different format
        # no clear speaker info
        timestamp = re.compile("[0-9:]+ Speaker")
        match = None
        for line in full_text:    
            if match is None: # keep track of when header info is complete
                match = re.match(timestamp, line.raw_text)
                line.useful_text = False
            elif re.match(timestamp, line.raw_text) is None and line.raw_text!="":
                line.text = unicodedata.normalize('NFKD',line.raw_text)
            else:
                line.useful_text = False
    elif focus_group:
        intro_end = raw_text.index('Transcriber: transcribername')
        full_text = [clean_line(match_speakers(line, focus_group)) if i>intro_end else line \
                    for i, line in enumerate(full_text) ]
    elif peer_healing:
        full_text = [clean_line(match_speakers(line, focus_group, peer_healing)) for line in full_text]
    else:
        full_text = [clean_line(match_speakers(line)) for line in full_text]
  
    filename = filepath.split('/')[-1]
    return InterviewFile(full_text, filename, focus_group)

# create datasets for modeling

class InterviewProc():
    '''Class for processing the interviews into chunks for modeling.
    Returns: list of dicts, each dict corresponds to a dataset item
    '''
    def __init__(self, filepaths):
        self.filepaths = filepaths
        # intiate empty values
        self.interviews = [None] * len(filepaths) # interview files read in at once

    def _basic_clean(self, text: str):
            # a modification from clean_data.py so that speaker name is kept
            dirtytext = re.compile(r"(\(.+?\))|(\*\*)|(\[-*\])") # match text in (text), **, or [--]
            clean_text = re.sub(dirtytext, "", text).strip()
            clean_text = clean_text.replace(u'\xa0', u' ')
            return clean_text
    
    def _chunk_text(self, lines:list,  max_count = 256, chunk_overlap = 0):
        '''Given a list of text, chunk each based on indent or sentences
        Cleans out any hanging punctuation
        Parameters:
            lines: list of text to be chunked
            max_count: max number of tokens
            chunk_overlap: number of tokens to overlap by
        Returns:
            list of chunks
        '''
        count_words = lambda x: len(x.split())
        text_splitter = RecursiveCharacterTextSplitter(chunk_size = max_count, 
                                                    chunk_overlap = chunk_overlap,
                                                    length_function =  count_words, 
                                                separators = ["\n\n", "\n", ".", "?", "!"])
        punct_trim = re.compile(r'^(\.|\?|!|\n)') #get rid of beginning punctuation
        chunks = text_splitter.create_documents(lines)
        clean_chunks = [re.sub(punct_trim, '', chunk.page_content).strip() for chunk in chunks]
        return clean_chunks
    
    def read_files(self):
        # read interviews from list of filepaths
        for i,f in enumerate(self.filepaths):
            with open(f, 'rb') as f:
                self.interviews[i] = pickle.load(f)
        
    def get_full_text(self):
        '''For passing full interview into the model
        Uses list of InterviewFile objects
        Returns:
           data: dict, filename and list of each interview's cleaned text
        ''' 
        data = {}
        for interview in self.interviews:
            clean_lines = [interview.full_text[j] for j in interview.clean_lines] 
            lines = []
            for line in clean_lines:
                if line.interviewee:
                    speaker = "Subject: "
                elif line.interviewer:
                    speaker = "Interviewer: "
                else: 
                    speaker = "Other: "
                lines.append(f'{speaker}{line.text}')
            data[interview.filename] = '\n '.join(lines)                
        return data

    def get_chunks(self, method = "paired", **kwargs):
        '''Extract list of chunks 
        Parameters:
            self.interviews: list of InterviewFile objects 
            method: str
                paired: interviewer appended to only the subsequent respondant statement 
                interviewee: respondant + any unknown speakers
                interviewer: interviewer and credible messengers
                other: returns all lines 
            **kwargs: optional parameters for `self._get_chunks`
                max_count, chunk_overlap
        Returns:
            data: dict, enumeration of chunked output across all interviews
        '''
      
        if method == "all":
            datafull = self.get_full_text()
            data = {f'{key}_{i}': chunk for key, text in datafull.items() 
                    for i, chunk in enumerate( self._chunk_text([text], **kwargs))}
        else:
            data = {}
            for _, interview in enumerate(self.interviews):
                data_list = []
                clean_lines = [interview.full_text[j] for j in interview.clean_lines] 
                if method =="paired":
                    skip_loop = False
                    for num, line in enumerate(clean_lines):
                        if skip_loop: # line was already saved 
                            skip_loop = False
                            continue
                        if num+1 < len(clean_lines):
                            if line.interviewer and clean_lines[num+1].interviewee:
                                # combine the two lines
                                combo_string = line.text + ' ' + clean_lines[num+1].text
                                data_list.append( combo_string)
                                skip_loop = True # skip the next loop since we appended it
                            else: # don't combine next line, don't skip next loop
                                data_list.append(line.text)
                                skip_loop = False
                        else: # it's the last line
                            data_list.append(line.text)
                elif method == 'interviewee':
                    data_list += [line.text for line in clean_lines if \
                    (line.credible_messenger is False) and (line.interviewer is False)]
                elif method == "interviewer":
                    data_list += [line.text for line in clean_lines if \
                                line.interviewer or line.credible_messenger]
                # chunk them in case theyre too long
                data.update({f'{interview.filename}_{i}': chunk for i,
                        chunk in enumerate(self._chunk_text(data_list, **kwargs))})
        return data

    def get_question_chunks(self, protocol_file, threshold = .2):
        '''Pass interview by questions.
        Parameters:
            protocol_file: filepath for document of question protocols
            threshold: similarity threshold
        Returns:
            data: dict of questions: responses combined
        '''
        from docx import Document
        from sentence_transformers import SentenceTransformer, util, losses
        intervieweetext = [line.text for interview in self.interviews for line in interview.full_text if \
                     (line.useful_text) and \
                     (line.credible_messenger is False) and \
                     (line.interviewer is False)]
        doc = Document(protocol_file)
        # get protocol questions
        protocolQuestions = []
        for pg in doc.paragraphs:
            text = pg.text.strip()
            # Only the lines ending with '?' get appended to the list, so the headers like "Stigma" won't get appended
            if (text.endswith("?")):
                protocolQuestions.append(text)
        protocolQuestions = tuple(protocolQuestions)
        model1 = SentenceTransformer('all-mpnet-base-v2')
        model2 = SentenceTransformer('multi-qa-MiniLM-L6-cos-v1')
        # Question Embeddings
        qEmbed1 = model1.encode(protocolQuestions, convert_to_tensor = True)
        qEmbed2 = model2.encode(protocolQuestions, convert_to_tensor = True)
        qEmbed = torch.cat((qEmbed1, qEmbed2), dim=1)
        # Sentence Embeddings
        sEmbed1 = model1.encode(intervieweetext, convert_to_tensor = True)
        sEmbed2 = model2.encode(intervieweetext, convert_to_tensor = True)
        sEmbed = torch.cat((sEmbed1, sEmbed2), dim=1)
        # Looping through each sentence embedding
        # save questions to dict based on question indices
        data = {question :[] for question in protocolQuestions}
        data['other'] = []
        for index, sent in enumerate(sEmbed):
            # Score for how semantically similar the sentence is with the interview protocol questions
            similar = util.cos_sim(sent, qEmbed)
            highestScore = similar.max().item()
            if highestScore > threshold:
                # Index of the most similar question
                bestMatch = torch.argmax(similar)
                matchedQ = protocolQuestions[bestMatch]
                data[matchedQ].append(intervieweetext[index])
            else:
                data['other'].append(intervieweetext[index])
        data = {k: '\n '.join(v) for k,v in data.items()}
        return data
    
class InterviewDataset(Dataset):
    '''Pass in Processed Interview and convert to torch dataset
        Useful for batch processing (probably not necessary)
    '''
    def __init__(self, data_dict, qchunk = False):
        '''
        data_dict: dict, where each value is a unique item
        qchunk: bool, True if using chunked questions (data dict values are lists)
    
        '''
        if qchunk:
            # explode out questions and data
            if type(list(data_dict.values())[0]) is list:
                self.questions = [k if k!= "other" else "Miscellaneous questions" for k, v in data_dict.items() for i in range(len(v))]  
                self.data = [chunk for qset in list(data_dict.values()) for chunk in qset]
                self.filenames = [f"{i}__{k}" if k!= "other" else "Miscellaneous questions"
                    for k, v in data_dict.items() for i in range(len(v))]   #set filename to be the questions_number
            else: #full questions
                self.questions = [k if k!= "other" else "Miscellaneous questions" for k, v in data_dict.items()]  
                self.data =  list(data_dict.values())
                self.filenames = self.questions

        else:
            self.questions = [np.nan for _ in range(len(data_dict))]
            self.data = list(data_dict.values())  
            self.filenames = list(data_dict.keys())   

    def __len__(self):
        return len(self.data)
    
    def __getitem__(self, idx):
        return  {'INTERVIEW': self.data[idx], 
                     'QUESTION': self.questions[idx],
                     'filename': self.filenames[idx]}
    
